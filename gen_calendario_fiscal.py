#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador de Calendario Fiscal por pais (Excel + PDF + .ics)
Reutiliza el motor util_fiscal.py y las bases obligaciones_*.json.

Produce, para un pais y anio dado:
  - calendario_fiscal_<PAIS>_<ANIO>.xlsx   (una hoja por mes, formato condicional)
  - calendario_fiscal_<PAIS>_<ANIO>.pdf    (version imprimible resumida)
  - calendario_fiscal_<PAIS>_<ANIO>.ics    (para Google Calendar / Outlook)

Uso:
  py -3.14 gen_calendario_fiscal.py --pais SV --anio 2027
  py -3.14 gen_calendario_fiscal.py --pais CO --anio 2027 --outdir "C:/salida"
"""
import os, sys, argparse, json, glob, calendar
from datetime import date, datetime, timedelta

# Permitir importar util_fiscal desde este script o desde cualquier cwd
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import util_fiscal

def obtener_datos_pais(codigo):
    """Carga la base JSON de un pais y la entidad fiscal."""
    return util_fiscal.cargar_pais(codigo)

# ---- Mapa de autoridades por pais (para el PDF y encabezados) ----
ENTIDADES = {
    "SV":"Ministerio de Hacienda (MH)","GT":"Superintendencia de Administración Tributaria (SAT)",
    "HN":"Servicio de Administración de Rentas (SAR)","NI":"Dirección General de Ingresos (DGI)",
    "PA":"Dirección General de Ingresos (DGI-MEF)","CR":"Ministerio de Hacienda (DGT)",
    "DO":"Dirección General de Impuestos Internos (DGII)","MX":"Servicio de Administración Tributaria (SAT)",
    "CO":"Dirección de Impuestos y Aduanas Nacionales (DIAN)","PE":"Superintendencia Nacional de Aduanas (SUNAT)",
    "CL":"Servicio de Impuestos Internos (SII)","AR":"ARCA","EC":"Servicio de Rentas Internas (SRI)",
    "UY":"Dirección General Impositiva (DGI)","PY":"Subsecretaría de Estado de Tributación (SET)",
    "BO":"Servicio de Impuestos Nacionales (SIN)","ES":"Agencia Estatal de Administración Tributaria (AEAT)",
    "GQ":"Ministerio de Hacienda","US":"Internal Revenue Service (IRS)",
}

def _siguiente_habil(d):
    """Avanza a un dia que no sea sabado/domingo (saltos fin de semana)."""
    while d.weekday() >= 5:
        d = d + timedelta(days=1)
    return d

def vencimientos_mes(data, anio, mes, digito=None):
    """Devuelve lista de (fecha, obligacion) que vencen en el mes dado.
    Si 'digito' se indica y la obligacion tiene dias_por_digito, calcula la
    fecha exacta de ese digito (con correccion a dia habil)."""
    hoy = date(anio, mes, 1)
    ultimo = calendar.monthrange(anio, mes)[1]
    res = util_fiscal.calcular_vencimientos_ventana(data, hoy, ultimo + 2)
    regla = data.get('reglas_vencimiento', {})
    dias_por_digito = regla.get('dias_por_digito') or {}
    out = []
    for r in res:
        f = date.fromisoformat(r['fecha'])
        oid = r['id']
        # Si la obligacion es escalonada por digito y estamos en modo digito,
        # recalcular la fecha con el dia real de ese digito.
        if digito is not None and oid in dias_por_digito:
            tabla = dias_por_digito[oid]
            if str(digito) in tabla:
                dia_d = tabla[str(digito)]
                try:
                    # el dia del digito aplica al mismo mes de la fecha base
                    f = _siguiente_habil(date(f.year, f.month, dia_d))
                except ValueError:
                    pass
        if f.year == anio and f.month == mes:
            out.append((f, r))
    return out

def generar_xlsx(data, anio, ciudad, outdir, digito=None):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    pais = data['pais']
    wb = Workbook()
    # Etiqueta de modo digito en encabezado
    modo = f"  ·  Dígito {digito}" if digito is not None else ""
    # Hoja resumen
    ws = wb.active
    ws.title = "Resumen"
    ws['A1'] = f"Calendario Fiscal {anio} · {pais}{modo}"
    ws['A1'].font = Font(bold=True, size=14)
    ws['A2'] = f"Fuente: {ENTIDADES.get(data['codigo'], 'autoridad fiscal')}"
    ws['A3'] = f"Actualizado: {data.get('ultima_actualizacion','')}"
    ws['A4'] = "Una hoja por mes con los vencimientos de cada declaración. Vencidos en rojo, próximos en amarillo."

    # Cabecera tablas mensuales
    headers = ["Fecha", "Formulario", "Obligación", "Tipo", "Estado"]
    fill = PatternFill("solid", fgColor="0E7C7B")
    hfont = Font(bold=True, color="FFFFFF")

    for mes in range(1, 13):
        s = wb.create_sheet(title=calendar.month_abbr[mes])
        s.append(headers)
        for c in range(1, len(headers)+1):
            cell = s.cell(row=1, column=c)
            cell.fill = fill; cell.font = hfont
        vencs = vencimientos_mes(data, anio, mes, digito)
        for f, r in sorted(vencs, key=lambda x: x[0]):
            hoy = date.today()
            estado = "Vencido" if f < hoy else ("Proximo" if f <= hoy + timedelta(days=7) else "En tiempo")
            s.append([f.isoformat(), r.get('formulario',''), r.get('nombre',''), r.get('tipo',''), estado])
        # Formato condicional (rojo/amarillo)
        from openpyxl.formatting.rule import CellIsRule
        last = s.max_row
        rng = f"E2:E{max(last,2)}"
        if last > 1:
            s.conditional_formatting.add(rng, CellIsRule(operator='equal', formula=['"Vencido"'], fill=PatternFill("solid", fgColor="F8696B")))
            s.conditional_formatting.add(rng, CellIsRule(operator='equal', formula=['"Proximo"'], fill=PatternFill("solid", fgColor="FFEB91")))
        # Acabado profesional (skill xlsx): congelar encabezado + autofiltro + anchos
        s.freeze_panes = "A2"
        s.auto_filter.ref = f"A1:E{max(last,1)}"
        s.row_dimensions[1].height = 20
        for c in range(1, len(headers)+1):
            s.column_dimensions[get_column_letter(c)].width = 14 if c in (1,2,4) else 42
        s.column_dimensions["E"].width = 12

    out = os.path.join(outdir, f"calendario_fiscal_{data['codigo'].lower()}_{anio}.xlsx")
    wb.save(out)
    return out

def generar_ics(data, anio, outdir, digito=None):
    pais = data['pais']
    import datetime as dt
    lines = [
        "BEGIN:VCALENDAR", "VERSION:2.0", "PRODID:-//Calendario Fiscal//ES",
        "CALSCALE:GREGORIAN", f"X-WR-CALNAME:Calendario Fiscal {anio} {pais}" + (f" (digito {digito})" if digito is not None else ""),
    ]
    seen = set()
    for mes in range(1, 13):
        for f, r in vencimientos_mes(data, anio, mes, digito):
            key = (f.isoformat(), r.get('id'))
            if key in seen: continue
            seen.add(key)
            uid = f"{r.get('formulario','')}-{f.isoformat()}-{data['codigo']}"
            lines += [
                "BEGIN:VEVENT",
                f"UID:{uid}@calendariofiscal",
                f"DTSTAMP:{dt.datetime.now(dt.timezone.utc).strftime('%Y%m%dT%H%M%SZ')}",
                f"DTSTART;VALUE=DATE:{f.strftime('%Y%m%d')}",
                f"DTEND;VALUE=DATE:{(f+timedelta(days=1)).strftime('%Y%m%d')}",
                f"SUMMARY:{r.get('formulario','')} - {r.get('nombre','')}",
                f"DESCRIPTION:Vence en la fecha indicada. Obligacion fiscal {pais}.",
                "END:VEVENT",
            ]
    lines.append("END:VCALENDAR")
    out = os.path.join(outdir, f"calendario_fiscal_{data['codigo'].lower()}_{anio}.ics")
    with open(out, "w", encoding="utf-8") as fh:
        fh.write("\r\n".join(lines))
    return out

def generar_pdf(data, anio, outdir, digito=None):
    from docx import Document
    from docx.shared import Pt, RGBColor
    pais = data['pais']
    doc = Document()
    st = doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10)
    p = doc.add_paragraph()
    modo = f"  ·  Dígito {digito}" if digito is not None else ""
    r = p.add_run(f"Calendario Fiscal {anio} · {pais}{modo}")
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x0E,0x7C,0x7B)
    doc.add_paragraph(f"Fuente: {ENTIDADES.get(data['codigo'],'autoridad fiscal')}   ·   Actualizado: {data.get('ultima_actualizacion','')}")
    doc.add_paragraph()
    for mes in range(1, 13):
        doc.add_paragraph()
        hr = doc.add_paragraph(); hm = hr.add_run(calendar.month_name[mes]); hm.bold=True; hm.font.size=Pt(12); hm.font.color.rgb=RGBColor(0x23,0x36,0x37)
        vencs = vencimientos_mes(data, anio, mes, digito)
        if not vencs:
            doc.add_paragraph("(sin vencimientos este mes)")
        for f, r in sorted(vencs, key=lambda x: x[0]):
            doc.add_paragraph(f"{f.isoformat()}  ·  {r.get('formulario','')}  ·  {r.get('nombre','')}")
    out = os.path.join(outdir, f"calendario_fiscal_{data['codigo'].lower()}_{anio}.pdf")
    doc.save(out)
    return out

if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--pais", required=True, help="Codigo de pais (SV, CO, MX...)")
    ap.add_argument("--anio", type=int, default=2027)
    ap.add_argument("--digito", type=int, default=None, help="Terminacion de NIT/CUIT (0-9) para calendarios por digito (CO, AR)")
    ap.add_argument("--outdir", default=os.path.join(os.environ.get('USERPROFILE','.'), 'CalendariosFiscales'))
    a = ap.parse_args()
    os.makedirs(a.outdir, exist_ok=True)
    data = obtener_datos_pais(a.pais)
    if not data:
        print(f"❌ Pais '{a.pais}' no encontrado. Disponibles: {', '.join(c for c,_ in util_fiscal.listar_paises())}")
        sys.exit(1)
    xlsx = generar_xlsx(data, a.anio, "", a.outdir, a.digito)
    ics = generar_ics(data, a.anio, a.outdir, a.digito)
    pdf = generar_pdf(data, a.anio, a.outdir, a.digito)
    print("✅ Generados:")
    print(f"   Excel: {xlsx}")
    print(f"   PDF  : {pdf}")
    print(f"   ICS  : {ics}")
